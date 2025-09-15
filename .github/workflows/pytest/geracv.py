from docx import Document

doc = Document()
doc.add_heading('Currículo de Jezuino Motta', 0)

doc.add_heading('📍 Dados Pessoais', level=1)
doc.add_paragraph('Nome: Jezuino Motta\nEmail: mottajezuino@gmail.com\nCidade: Peruíbe - SP')

doc.add_heading('🎓 Formação Acadêmica', level=1)
doc.add_paragraph('Bacharel em Ciência da Computação – Universidade Federal de São Paulo (2018–2022)')

doc.add_heading('💼 Experiência Profissional', level=1)
doc.add_paragraph('Desenvolvedor Python – TechNova (2022–2024)\n- Desenvolvimento de aplicações web com Flask e Streamlit\n- Integração de APIs e automação de processos com Python')

doc.add_heading('🧠 Habilidades Técnicas', level=1)
doc.add_paragraph('- Python, Streamlit, Flask\n- Git, Docker, Linux\n- Google Gemini API\n- Análise de dados com Pandas e NumPy')

doc.add_heading('🌐 Idiomas', level=1)
doc.add_paragraph('- Português: Nativo\n- Inglês: Intermediário')

doc.save('pytest/exemplo.docx')
